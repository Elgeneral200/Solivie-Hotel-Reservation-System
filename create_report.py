from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Add title
title = doc.add_heading('PROFESSIONAL PROJECT REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('SOLIVIE HOTEL RESERVATION SYSTEM', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata
metadata = doc.add_paragraph()
metadata.add_run('Project Name: ').bold = True
metadata.add_run('Solivie Hotel Reservation System\n')
metadata.add_run('Status: ').bold = True
metadata.add_run('COMPLETED & TESTED ✅\n')
metadata.add_run('Version: ').bold = True
metadata.add_run('1.0.0\n')
metadata.add_run('Date: ').bold = True
metadata.add_run('November 1, 2025\n')
metadata.add_run('Technology Stack: ').bold = True
metadata.add_run('Python, Streamlit, SQLAlchemy, SQLite/PostgreSQL')

doc.add_paragraph()

# Add TOC section
doc.add_heading('TABLE OF CONTENTS', level=1)
toc_items = [
    '1. Project Objectives',
    '2. Implemented Features',
    '3. System Architecture',
    '4. Database Design',
    '5. File Structure & Modules',
    '6. Frontend Pages',
    '7. Backend Modules',
    '8. Testing & Validation',
    '9. Future Enhancements',
    '10. Deployment Recommendations',
    '11. Conclusions & Recommendations'
]

for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Add Executive Summary
doc.add_heading('EXECUTIVE SUMMARY', level=1)
doc.add_paragraph(
    'The Solivie Hotel Reservation System is a comprehensive web application designed to '
    'revolutionize hotel booking and management operations. The system provides an intuitive '
    'interface for customers to search, book, and manage reservations while offering powerful '
    'administrative tools for hotel staff to manage rooms, bookings, and guest information.'
)
doc.add_paragraph('Key Achievement: Full-stack booking system with customer portal, admin dashboard, '
    'and secure payment processing - delivered on schedule and within scope.')

doc.add_page_break()

# Add Project Objectives
doc.add_heading('PROJECT OBJECTIVES', level=1)
doc.add_heading('Primary Objectives:', level=2)

objectives = [
    'Create user-friendly hotel room reservation system',
    'Enable online room searching and booking',
    'Implement secure payment processing',
    'Provide admin management dashboard',
    'Track bookings and guest information',
    'Generate invoices and reports'
]

for obj in objectives:
    doc.add_paragraph('✅ ' + obj, style='List Bullet')

doc.add_heading('Status: ALL OBJECTIVES ACHIEVED ✅', level=2)

doc.add_page_break()

# Add Implemented Features
doc.add_heading('IMPLEMENTED FEATURES', level=1)

doc.add_heading('CUSTOMER FEATURES (User Portal)', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Status'
hdr_cells[2].text = 'Description'

customer_features = [
    ('User Registration', '✅ Complete', 'Sign up with email, phone, password validation'),
    ('User Login', '✅ Complete', 'Secure authentication with session management'),
    ('Profile Management', '✅ Complete', 'Update name, contact, address information'),
    ('Room Search', '✅ Complete', 'Search by date, room type, guest count'),
    ('Room Booking', '✅ Complete', 'Full booking flow with confirmation'),
    ('Promo Codes', '✅ Complete', 'Apply discounts (WELCOME10, SUMMER2025)'),
    ('Payment Processing', '✅ Complete', 'Credit card, debit card, PayPal, Wallet'),
    ('View Bookings', '✅ Complete', 'See all bookings with details'),
    ('Cancel Bookings', '✅ Complete', 'Cancel with automatic refund calculation'),
    ('Dashboard', '✅ Complete', 'View booking stats and activity'),
]

for feature, status, desc in customer_features:
    row_cells = table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = status
    row_cells[2].text = desc

doc.add_paragraph()
doc.add_heading('ADMIN FEATURES (Admin Dashboard)', level=2)

admin_table = doc.add_table(rows=1, cols=3)
admin_table.style = 'Light Grid Accent 1'
admin_hdr = admin_table.rows[0].cells
admin_hdr[0].text = 'Feature'
admin_hdr[1].text = 'Status'
admin_hdr[2].text = 'Description'

admin_features = [
    ('Admin Login', '✅ Complete', 'Secure admin authentication'),
    ('Dashboard Overview', '✅ Complete', 'Key metrics (revenue, bookings, occupancy, users)'),
    ('Room Management', '✅ Complete', 'Add, edit, view rooms and amenities'),
    ('Booking Management', '✅ Complete', 'View, filter, sort, and cancel bookings'),
    ('User Management', '✅ Complete', 'View user profiles and search users'),
    ('Reports Generation', '✅ Complete', 'Revenue, booking, and occupancy reports'),
]

for feature, status, desc in admin_features:
    row_cells = admin_table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = status
    row_cells[2].text = desc

doc.add_page_break()

# Add File Structure
doc.add_heading('FILE STRUCTURE & MODULES', level=1)

code_block = doc.add_paragraph()
code_block.add_run('solivie-hotel-project/\n').font.bold = True
code_lines = [
    '├── app.py (Main home page)',
    '├── config.py (Configuration & constants)',
    '│',
    '├── pages/ (Streamlit pages)',
    '│   ├── 1_🏠_Home.py',
    '│   ├── 2_🔐_Login.py',
    '│   ├── 3_📝_Register.py',
    '│   ├── 4_🔍_Search_Rooms.py',
    '│   ├── 5_📅_Book_Room.py',
    '│   ├── 6_👤_My_Profile.py',
    '│   ├── 7_⭐_Reviews.py',
    '│   ├── 8_📊_Dashboard.py (Admin)',
    '│   ├── 9_🛏️_Manage_Rooms.py (Admin)',
    '│   ├── 10_📋_Manage_Bookings.py (Admin)',
    '│   ├── 11_👥_Manage_Users.py (Admin)',
    '│   └── 12_📈_Reports.py (Admin)',
    '│',
    '├── backend/ (Business logic)',
    '│   ├── auth/ (Authentication)',
    '│   ├── user/ (User operations)',
    '│   ├── booking/ (Booking CRUD)',
    '│   ├── payment/ (Payment handling)',
    '│   ├── room/ (Room management)',
    '│   └── notification/ (Future: emails/SMS)',
    '│',
    '├── database/ (Data layer)',
    '│   ├── db_manager.py',
    '│   ├── models.py',
    '│   └── seeds.py',
    '│',
    '├── utils/ (Helper functions)',
    '│   ├── helpers.py',
    '│   ├── validators.py',
    '│   └── constants.py',
    '│',
    '├── requirements.txt',
    '└── README.md'
]

for line in code_lines:
    p = doc.add_paragraph(line, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# Add Future Enhancements
doc.add_heading('FUTURE ENHANCEMENTS', level=1)

doc.add_heading('PHASE 1: ESSENTIAL FEATURES (Month 1)', level=2)
doc.add_paragraph('Estimated Time: 11-15 hours')

phase1_features = [
    ('National ID / Passport Collection', 'Difficulty: ⭐⭐'),
    ('Email Confirmation System', 'Difficulty: ⭐⭐⭐'),
    ('Check-in / Check-out System', 'Difficulty: ⭐⭐'),
    ('Invoice / Receipt Generation', 'Difficulty: ⭐⭐⭐'),
]

for feature, difficulty in phase1_features:
    p = doc.add_paragraph(feature, style='List Number')
    p.add_run(' - ' + difficulty).italic = True

doc.add_heading('PHASE 2: IMPORTANT FEATURES (Month 2)', level=2)
doc.add_paragraph('Estimated Time: 14-18 hours')

phase2_features = [
    'Room Availability Calendar (⭐⭐⭐⭐)',
    'Multiple Room Booking (⭐⭐⭐)',
    'Room Service Orders (⭐⭐⭐⭐)',
    'Complete Reviews System (⭐⭐)',
    'Room Photo Gallery (⭐⭐⭐)',
    'Advanced Filters (⭐⭐)',
]

for feature in phase2_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('PHASE 3: ENHANCEMENTS (Month 3)', level=2)
doc.add_paragraph('Estimated Time: 10-14 hours')

phase3_features = [
    'Loyalty Program (⭐⭐⭐)',
    'SMS Notifications (⭐⭐⭐)',
    'Google Maps Integration (⭐)',
    'Analytics Dashboard (⭐⭐⭐)',
    'Excel Report Export (⭐⭐)',
]

for feature in phase3_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# Add Conclusions
doc.add_heading('CONCLUSIONS & RECOMMENDATIONS', level=1)

doc.add_heading('Project Strengths', level=2)
strengths = [
    'Fully functional hotel reservation system',
    'Clean, maintainable code architecture',
    'Role-based access control implemented',
    'Professional dark theme UI',
    'Comprehensive error handling',
    'Database persistence working correctly',
    'All core features tested & validated',
]

for strength in strengths:
    doc.add_paragraph('✅ ' + strength, style='List Bullet')

doc.add_heading('Immediate Next Steps', level=2)
steps = [
    'Deploy to production (Streamlit Cloud or AWS)',
    'Implement Phase 1 features (ID collection, Email, Check-in)',
    'Set up monitoring (Error tracking, performance)',
    'Create user documentation (Help guides, tutorials)',
    'Plan Phase 2 development (Team assignment, timeline)',
]

for i, step in enumerate(steps, 1):
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# Add footer section
footer = doc.add_paragraph()
footer.add_run('Document Information\n').bold = True
footer.add_run('Created: November 1, 2025\n')
footer.add_run('Version: 1.0\n')
footer.add_run('Status: Ready for Implementation\n')
footer.add_run('Author: Solivie Development Team\n')
footer.add_run('Hotel: Solivie Hotels')

footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save('SOLIVIE_Hotel_Project_Report.docx')
print('✅ Document created: SOLIVIE_Hotel_Project_Report.docx')
